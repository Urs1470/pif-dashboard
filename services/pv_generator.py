"""
Generator PV (proces verbal) din template-uri DOCX.

Strategie:
- Header / footer / stilurile documentului raman din template (sigla firmei, marginile, fonturile).
- Continutul (paragrafe + tabele) e reconstruit programatic pentru a garanta:
  * format consistent (data DD.MM, fara wrap urat)
  * tabele aliniate cu chenare drepte
  * semnaturi cu firma aliniata vertical cu "(Numele, prenumele si semnatura)"
  * Observatii: forteaza page break inainte (mereu pe pagina 2)
"""
from copy import deepcopy
from datetime import datetime
import os
import re
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.table import Table

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'templates', 'pv')
SERVICE_TEMPLATE = os.path.join(TEMPLATE_DIR, 'PV_SERVICE_TEMPLATE.docx')
PIF_TEMPLATE = os.path.join(TEMPLATE_DIR, 'PV_PIF_250592E_C1.docx')

LEGEND_ACTIUNI = [
    ('TD',  'TIMP DEPLASARE'),
    ('TP',  'timp procesare'),
    ('TDI', 'timp deplasare intoarcere'),
    ('TPP', 'timp postprocesare'),
    ('OLS', 'ore lucrate service'),
    ('TA',  'Timp asteptare'),
    ('CZA', 'cazare'),
]


# ----------- helpers paragrafe -----------

def _set_paragraph_text(paragraph, new_text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def _para_text(paragraph):
    return ''.join(r.text for r in paragraph.runs)


def _replace_in_paragraph_prefix(paragraph, prefix, new_value_after_prefix):
    txt = _para_text(paragraph)
    if txt.strip().startswith(prefix):
        _set_paragraph_text(paragraph, f"{prefix}{new_value_after_prefix}")
        return True
    return False


def _delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def _add_page_break_before(paragraph):
    """Forteaza paragraful sa inceapa pe pagina noua."""
    pPr = paragraph._p.get_or_add_pPr()
    pageBreakBefore = OxmlElement('w:pageBreakBefore')
    pPr.append(pageBreakBefore)


# ----------- helpers tabele -----------

def _norm_data(d):
    """Normalizeaza data la format DD.MM (taie anul daca apare)."""
    if not d:
        return ''
    m = re.match(r'(\d{1,2})[./](\d{1,2})', d.strip())
    if m:
        return f"{m.group(1).zfill(2)}.{m.group(2).zfill(2)}"
    return d.strip()


def _apply_grid_borders(table, sz='4', color='000000'):
    """Aplica chenare drepte pe tot tabelul (substitut pentru style='Table Grid')."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # remove existing tblBorders
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_cell_borders(cell, sz='4', color='000000', sides=('top','left','bottom','right')):
    tcPr = cell._tc.get_or_add_tcPr()
    # remove existing borders
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in sides:
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    # for edges not in sides, set "nil" (no border)
    all_edges = {'top','left','bottom','right'}
    for edge in all_edges - set(sides):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill='F2F2F2'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, italic=False, size=10,
                   align=None, valign=None, font_name='Calibri'):
    cell.text = ''  # reset
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if valign is not None:
        cell.vertical_alignment = valign


def _set_cell_width(cell, cm_width):
    cell.width = Cm(cm_width)
    # also set in tcW
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(cm_width * 567)))  # 567 twips per cm
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _vertical_merge_restart(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    vMerge.set(qn('w:val'), 'restart')
    tcPr.append(vMerge)


def _vertical_merge_continue(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    tcPr.append(vMerge)


def _remove_table(tbl):
    """Sterge un tabel din document. Returneaza elementul XML parent + index pentru insertie."""
    tbl_elem = tbl._tbl
    parent = tbl_elem.getparent()
    idx = list(parent).index(tbl_elem)
    parent.remove(tbl_elem)
    return parent, idx


def _move_table_to_position(table, parent, idx):
    """Muta tabelul la pozitia (parent, idx) in document."""
    tbl_elem = table._tbl
    cur_parent = tbl_elem.getparent()
    if cur_parent is not None:
        cur_parent.remove(tbl_elem)
    parent.insert(idx, tbl_elem)


def _find_table_with_text(doc, needle):
    for tbl in doc.tables:
        for row in tbl.rows:
            joined = ' | '.join(c.text for c in row.cells).lower()
            if needle.lower() in joined:
                return tbl
    return None


def _find_paragraph_with_text(doc, needle_startswith):
    for p in doc.paragraphs:
        if _para_text(p).strip().startswith(needle_startswith):
            return p
    return None


# ----------- build tabel ore -----------

def _build_hours_table(doc, ore, luna_pv):
    """
    Construieste un tabel nou (la finalul docului) cu structura:
      Row 0: titlu "Ore lucrate / alte activitati / cheltuieli in luna MM.YYYY" (merge orizontal)
      Row 1: header — Tip | Data | Durata | Nr ing | Ore (merge 2 col) | Observatii
      Row 2: sub-header — (vmerge) | (vmerge) | (vmerge) | (vmerge) | Normal | Suplimentar | (vmerge)
      Row 3+: data rows
    7 coloane.
    """
    ncols = 7
    nrows = 3 + max(len(ore), 1)  # min 1 rand data ca sa nu ramana tabel cu doar header
    table = doc.add_table(rows=nrows, cols=ncols)
    table.autofit = False
    _apply_grid_borders(table)

    col_widths_cm = [2.3, 1.7, 2.8, 1.6, 1.8, 1.8, 4.0]

    # Row 0 — titlu
    row0 = table.rows[0]
    title_cell = row0.cells[0]
    for c in row0.cells[1:]:
        title_cell.merge(c)
    _set_cell_text(title_cell, f'Ore lucrate / alte activitati / cheltuieli in luna {luna_pv}',
                   bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
                   valign=WD_ALIGN_VERTICAL.CENTER)
    _set_cell_shading(title_cell, 'E8F1F2')

    # Row 1 — header
    row1 = table.rows[1]
    headers_r1 = ['Tip activitate', 'Data', 'Durata\nDe la - la', 'Nr. de ingineri', 'Ore', None, 'Observatii']
    for i, txt in enumerate(headers_r1):
        if txt is None:
            continue
        cell = row1.cells[i]
        # split pe linii daca e cazul
        cell.text = ''
        for li, line in enumerate(txt.split('\n')):
            if li == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, 'F2F2F2')

    # Merge orizontal "Ore" peste coloanele 4 si 5
    row1.cells[4].merge(row1.cells[5])

    # Row 2 — sub-header Normal/Suplimentar
    row2 = table.rows[2]
    _set_cell_text(row2.cells[4], 'Normal', bold=True, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER, valign=WD_ALIGN_VERTICAL.CENTER)
    _set_cell_shading(row2.cells[4], 'F2F2F2')
    _set_cell_text(row2.cells[5], 'Suplimentar', bold=True, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER, valign=WD_ALIGN_VERTICAL.CENTER)
    _set_cell_shading(row2.cells[5], 'F2F2F2')

    # Vertical merge: coloanele 0,1,2,3,6 din row1 + row2 sunt aceeasi celula
    for col in (0, 1, 2, 3, 6):
        _vertical_merge_restart(row1.cells[col])
        _vertical_merge_continue(row2.cells[col])

    # Data rows
    for idx, entry in enumerate(ore or [{}]):
        r = table.rows[3 + idx]
        values = [
            (entry.get('tip') or '').strip(),
            _norm_data(entry.get('data') or ''),
            (entry.get('durata') or '').strip(),
            str(entry.get('nr_ingineri') or 1),
            str(entry.get('ore_normal') or '').strip(),
            str(entry.get('ore_supl') or '').strip(),
            (entry.get('observatii') or '').strip(),
        ]
        for i, val in enumerate(values):
            _set_cell_text(r.cells[i], val, size=10,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           valign=WD_ALIGN_VERTICAL.CENTER)

    # Setare latimi coloane pe TOATE randurile
    for row in table.rows:
        for i, w in enumerate(col_widths_cm):
            if i < len(row.cells):
                _set_cell_width(row.cells[i], w)

    return table


def _build_legend_table(doc):
    """
    Tabel legenda 2 coloane × N randuri (cate 1 categorie per celula).
    Prima coloana are 4 categorii, a doua 3.
    Adaugam prima coloana cu eticheta "Legenda" merged peste tot.
    """
    left = LEGEND_ACTIUNI[:4]   # TD, TP, TDI, TPP
    right = LEGEND_ACTIUNI[4:]  # OLS, TA, CZA
    nrows = max(len(left), len(right))
    table = doc.add_table(rows=nrows, cols=3)
    table.autofit = False
    _apply_grid_borders(table)
    col_widths_cm = [2.0, 6.5, 6.5]

    # Coloana 0: "Legenda" merged
    legend_cell = table.rows[0].cells[0]
    for i in range(1, nrows):
        legend_cell.merge(table.rows[i].cells[0])
    _set_cell_text(legend_cell, 'Legenda', bold=True, size=10,
                   align=WD_ALIGN_PARAGRAPH.CENTER, valign=WD_ALIGN_VERTICAL.CENTER)
    _set_cell_shading(legend_cell, 'F2F2F2')

    for i in range(nrows):
        if i < len(left):
            k, v = left[i]
            _set_cell_text(table.rows[i].cells[1], f'{k} = {v}', size=10,
                           valign=WD_ALIGN_VERTICAL.CENTER)
        if i < len(right):
            k, v = right[i]
            _set_cell_text(table.rows[i].cells[2], f'{k} = {v}', size=10,
                           valign=WD_ALIGN_VERTICAL.CENTER)

    for row in table.rows:
        for i, w in enumerate(col_widths_cm):
            if i < len(row.cells):
                _set_cell_width(row.cells[i], w)

    return table


def _build_ingineri_table(doc, ingineri):
    """
    Tabel 2 randuri x 3 coloane: header Inginer 1/2/3 + nume pe coloanele
    corespunzatoare. `ingineri` poate fi string (backward compat — 1 nume)
    sau lista de pana la 3 nume.
    """
    if isinstance(ingineri, str):
        ingineri = [ingineri] if ingineri.strip() else []
    ingineri = [n.strip() for n in (ingineri or []) if n and n.strip()][:3]

    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    _apply_grid_borders(table)
    col_w = 5.0

    for i, h in enumerate(['Inginer 1', 'Inginer 2', 'Inginer 3']):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       valign=WD_ALIGN_VERTICAL.CENTER)
        _set_cell_shading(table.rows[0].cells[i], 'F2F2F2')

    for i in range(3):
        nume = ingineri[i] if i < len(ingineri) else ''
        _set_cell_text(table.rows[1].cells[i], nume, size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       valign=WD_ALIGN_VERTICAL.CENTER)

    for row in table.rows:
        for i in range(3):
            _set_cell_width(row.cells[i], col_w)

    return table


def _build_signature_table(doc, firm_left, firm_right):
    """
    Tabel 1 rand x 2 coloane, fara chenar.
    Fiecare celula: firma (bold), 3 linii goale, "(Numele, prenumele si semnatura)" centrat.
    Garanteaza alinierea verticala intre firma si "(Numele...)".
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    for cell_idx, firma in enumerate([firm_left, firm_right]):
        cell = table.rows[0].cells[cell_idx]
        cell.text = ''
        # Linia 1: firma
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(firma)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        # 3 linii goale pentru spatiu de semnatura
        for _ in range(3):
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Linia "(Numele, prenumele si semnatura)"
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run('(Numele, prenumele si semnatura)')
        r2.italic = True
        r2.font.size = Pt(9)
        r2.font.name = 'Calibri'
        # No borders
        _set_cell_borders(cell, sides=())
        _set_cell_width(cell, 8.0)

    return table


# ----------- main entry -----------

def generate_pv_service(proiect, echipament, form_data):
    """
    proiect: dict din tabela proiecte (client, locatie, service_before, service_after, ...)
    echipament: dict din echipamente (nume, model, serial_number) | None
    form_data: dict cu cheile:
        cod_proiect, data_pv (DD.MM.YYYY), denumire_comanda,
        observatii (str|''), nota_facturare (str|''),
        reprezentant_eg (default 'Ion Ursu'),
        ore: [{tip, data, durata, nr_ingineri, ore_normal, ore_supl, observatii}]
    Returneaza bytes DOCX.
    """
    doc = Document(SERVICE_TEMPLATE)

    client = (proiect.get('client') or '').strip()
    locatie = (proiect.get('locatie') or '').strip()
    motivul = (proiect.get('service_before') or '').strip()
    actiuni = (proiect.get('service_after') or '').strip()

    if echipament:
        echi_model = (echipament.get('model') or echipament.get('nume') or '').strip()
        serie = (echipament.get('serial_number') or '').strip()
    else:
        echi_model = (proiect.get('echipament_principal') or '').strip()
        serie = ''

    cod = form_data.get('cod_proiect', '').strip()
    data_pv = form_data.get('data_pv', datetime.now().strftime('%d.%m.%Y'))
    denumire = form_data.get('denumire_comanda', '').strip()
    observatii = (form_data.get('observatii') or '').strip()
    nota_fact = (form_data.get('nota_facturare') or '').strip()
    rep_eg = (form_data.get('reprezentant_eg') or 'Ion Ursu').strip()
    ingineri = form_data.get('ingineri') or ([rep_eg] if rep_eg else ['Ion Ursu'])
    ore = form_data.get('ore') or []

    nr_comanda_text = f"{cod} {denumire}" if denumire else cod
    luna_pv = _luna_din_ore(ore) or datetime.now().strftime('%m.%Y')

    # ---- Inlocuiri text in paragrafe ----
    state = {'in_motivul': False, 'in_actiuni': False, 'observatii_next': False}
    motivul_block = []
    actiuni_block = []
    observatii_header_para = None
    observatii_content_para = None
    nota_fact_para_list = []

    for p in list(doc.paragraphs):
        txt = _para_text(p).strip()

        if re.match(r'^Nr\.\s*\d+\s*/', txt):
            _set_paragraph_text(p, f"Nr. 1/{cod}/{data_pv}")
            continue
        if _replace_in_paragraph_prefix(p, 'CLIENT:', f' {client}'):
            continue
        if _replace_in_paragraph_prefix(p, 'LOCATIE:', f' {locatie}'):
            continue
        if (_replace_in_paragraph_prefix(p, 'NR. COMANDĂ:', f' {nr_comanda_text}') or
            _replace_in_paragraph_prefix(p, 'NR. COMANDA:', f' {nr_comanda_text}')):
            continue
        if _replace_in_paragraph_prefix(p, 'ECHIPAMENT:', f' {echi_model}'):
            continue
        if _replace_in_paragraph_prefix(p, 'SERIE:', f' {serie}'):
            continue

        if txt.startswith('Motivul solicit'):
            state['in_motivul'] = True
            continue
        if state['in_motivul']:
            if txt.startswith('Au fost efectuate'):
                state['in_motivul'] = False
                state['in_actiuni'] = True
                continue
            if txt:
                motivul_block.append(p)
            continue
        if state['in_actiuni']:
            if txt.startswith('Observații') or txt.startswith('Observatii'):
                state['in_actiuni'] = False
                observatii_header_para = p
                state['observatii_next'] = True
                continue
            if txt:
                actiuni_block.append(p)
            continue
        if state.get('observatii_next') and txt:
            observatii_content_para = p
            state['observatii_next'] = False
            continue

        if txt.startswith('Intervenția se va factura') or txt.startswith('Interventia se va factura'):
            nota_fact_para_list.append(p)
            continue

    motivul_paras = _apply_block(motivul_block, motivul.splitlines() if motivul else [])
    actiuni_paras = _apply_block(actiuni_block, actiuni.splitlines() if actiuni else [])

    # ---- Insereaza poze pe ancore (Service) ----
    # Anchor format: "motiv:N", "actiuni:N", "observatii", "final"
    # unde N = indexul liniei din motivul/actiuni. Pozele se insereaza dupa
    # paragraful exact ales de Ion in modal.
    images = form_data.get('images') or []
    images_groups = _group_images_ordered(images)

    def _resolve_service_anchor(anchor_key):
        """Returneaza paragraph-ul dupa care se insereaza pozele, sau None."""
        if anchor_key.startswith('motiv:'):
            try:
                idx = int(anchor_key.split(':', 1)[1])
                return motivul_paras[idx] if 0 <= idx < len(motivul_paras) else (motivul_paras[-1] if motivul_paras else None)
            except (ValueError, IndexError):
                return motivul_paras[-1] if motivul_paras else None
        if anchor_key.startswith('actiuni:'):
            try:
                idx = int(anchor_key.split(':', 1)[1])
                return actiuni_paras[idx] if 0 <= idx < len(actiuni_paras) else (actiuni_paras[-1] if actiuni_paras else None)
            except (ValueError, IndexError):
                return actiuni_paras[-1] if actiuni_paras else None
        return None  # 'observatii' / 'final' tratate separat mai jos

    # Insereaza pozele cu ancore in motiv/actiuni (in ordine inversa pe acelasi
    # paragraf nu e cazul — fiecare grup are anchor unic)
    deferred_groups = []  # (anchor_key, imgs) pentru observatii / final
    for anchor_key, imgs in images_groups:
        target = _resolve_service_anchor(anchor_key)
        if target is not None:
            _build_images_block(doc, imgs, after_paragraph=target)
        else:
            deferred_groups.append((anchor_key, imgs))

    # Observatii: page break inainte de "Observații:" header + setare continut
    if observatii_header_para is not None:
        _add_page_break_before(observatii_header_para)
        if observatii:
            if observatii_content_para is not None:
                _set_paragraph_text(observatii_content_para, observatii)
            else:
                # adauga un paragraph dupa header
                new_p = observatii_header_para._p.addnext(deepcopy(observatii_header_para._p))
                # set text
                from docx.text.paragraph import Paragraph as DocxParagraph
                wrap = DocxParagraph(new_p, observatii_header_para._parent)
                _set_paragraph_text(wrap, observatii)
        else:
            # Daca nu sunt observatii, las paragraph-ul de continut gol vizibil
            if observatii_content_para is not None:
                _set_paragraph_text(observatii_content_para, 'Nu sunt.')

        # Imagini cu anchor "observatii"
        for anchor_key, imgs in deferred_groups:
            if anchor_key == 'observatii':
                anchor = observatii_content_para if (observatii_content_para is not None and observatii_content_para._p is not None) else observatii_header_para
                if anchor is not None:
                    _build_images_block(doc, imgs, after_paragraph=anchor)

    # Nota facturare
    if nota_fact:
        for p in nota_fact_para_list:
            _set_paragraph_text(p, nota_fact)
    else:
        for p in nota_fact_para_list:
            _delete_paragraph(p)

    # ---- Tabele: sterge cele vechi si rebuildare ----
    # Ordinea in document: 1) Ore lucrate  2) Legenda+ingineri
    old_hours = _find_table_with_text(doc, 'tip activitate')
    old_legend = _find_table_with_text(doc, 'inginer 1')

    insert_targets = []  # (parent, idx) pentru re-insertie

    if old_hours is not None:
        parent, idx = _remove_table(old_hours)
        insert_targets.append(('hours', parent, idx))
    if old_legend is not None:
        parent, idx = _remove_table(old_legend)
        insert_targets.append(('legend_ing', parent, idx))

    # Build tabele noi la finalul docului, apoi le mut la pozitiile vechi
    hours_table = _build_hours_table(doc, ore, luna_pv)
    legend_table = _build_legend_table(doc)
    ingineri_table = _build_ingineri_table(doc, ingineri)

    # Spacing paragraph empty intre tabelul ore si legenda
    # Si intre legenda si ingineri
    # python-docx adauga paragrafe dupa fiecare add_table — ok.

    # Re-pozitionez in document
    # Daca am gasit pozitia vechiului tabel ore, mut hours_table acolo
    target_hours = next((t for t in insert_targets if t[0] == 'hours'), None)
    target_legend = next((t for t in insert_targets if t[0] == 'legend_ing'), None)

    if target_hours:
        _, parent, idx = target_hours
        _move_table_to_position(hours_table, parent, idx)
        # daca era si legenda dupa, re-index
        if target_legend:
            target_legend = (target_legend[0], target_legend[1], target_legend[2])

    if target_legend:
        _, parent, idx = target_legend
        # adjust idx daca am inserat ceva inainte
        # mut legend si ingineri unul dupa altul
        _move_table_to_position(legend_table, parent, idx)
        # apoi spacing paragraph + ingineri
        # gaseste pozitia legend dupa mutare
        legend_idx = list(parent).index(legend_table._tbl)
        # adauga un paragraf gol intre
        spacer = OxmlElement('w:p')
        parent.insert(legend_idx + 1, spacer)
        _move_table_to_position(ingineri_table, parent, legend_idx + 2)

    # ---- Semnaturi ----
    # Sterg paragrafele cu "SC ELECTROGLOBAL SA" si urmatorul "(Numele...)"
    # si inserez tabel borderless cu firme + linii goale + (Numele...)
    sig_paragraph = _find_paragraph_with_text(doc, 'SC ELECTROGLOBAL') or \
                    _find_paragraph_with_text(doc, '         SC ELECTROGLOBAL')
    if sig_paragraph is None:
        # Cauta in lista paragrafelor
        for p in doc.paragraphs:
            if 'ELECTROGLOBAL' in _para_text(p):
                sig_paragraph = p
                break

    if sig_paragraph is not None:
        sig_parent = sig_paragraph._p.getparent()
        sig_idx = list(sig_parent).index(sig_paragraph._p)

        # Construiesc tabelul de semnatura la finalul docului
        sig_table = _build_signature_table(doc, 'SC ELECTROGLOBAL SA', client or '')

        # Sterg paragraful original cu firme
        _delete_paragraph(sig_paragraph)

        # Caut si sterg paragraful cu "(Numele, prenumele si semnatura)" — exista de 2 ori repetat
        to_delete = []
        for p in list(doc.paragraphs):
            t = _para_text(p)
            if 'Numele, prenumele' in t or 'semn' in t.lower() and 'prenumele' in t.lower():
                to_delete.append(p)
                if len(to_delete) >= 1:
                    break
        for p in to_delete:
            _delete_paragraph(p)

        # Mut tabelul la pozitia veche
        _move_table_to_position(sig_table, sig_parent, sig_idx)

    # Anexa final: pozele cu anchor "final" se adauga la sfarsitul docului
    final_imgs = [imgs for (a, imgs) in deferred_groups if a == 'final']
    if final_imgs:
        hdr_p = doc.add_paragraph()
        _add_page_break_before(hdr_p)
        run = hdr_p.add_run('Anexe foto')
        run.bold = True
        run.font.size = Pt(13)
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _build_images_block(doc, final_imgs[0], after_paragraph=hdr_p)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _luna_din_ore(ore):
    for r in ore:
        d = (r.get('data') or '').strip()
        m = re.match(r'(\d{1,2})\.(\d{1,2})(?:\.(\d{2,4}))?', d)
        if m:
            mm = m.group(2).zfill(2)
            yy = m.group(3) or datetime.now().strftime('%Y')
            if len(yy) == 2:
                yy = '20' + yy
            return f"{mm}.{yy}"
    return None


def _apply_block(paragraphs, lines):
    """Aplica `lines` pe `paragraphs`. Returneaza lista de paragrafe rezultate
    (in ordine, cate unul pe linie) — utila ca ancore pentru poze."""
    lines = [l.strip() for l in lines if l.strip()]
    result = []
    survivors = list(paragraphs)
    for i, p in enumerate(paragraphs):
        if i < len(lines):
            _set_paragraph_text(p, lines[i])
            result.append(p)
        else:
            _delete_paragraph(p)
            survivors[i] = None
    if len(lines) > len(paragraphs) and paragraphs:
        anchors = [p for p in survivors if p is not None and p._p is not None]
        anchor = anchors[-1] if anchors else None
        if anchor is None:
            return result
        from docx.text.paragraph import Paragraph as DocxParagraph
        extra = lines[len(paragraphs):]
        for line in extra:
            new_p = deepcopy(anchor._p)
            for t in new_p.iter(qn('w:t')):
                t.text = ''
            anchor._p.addnext(new_p)
            wrap = DocxParagraph(new_p, anchor._parent)
            _set_paragraph_text(wrap, line)
            anchor = wrap
            result.append(wrap)
    return result


def _format_date_list(dates):
    """['01.10.2026','22.10.2026'] -> '01.10.2026 si 22.10.2026'; lista lunga: ',' + ultima 'si'."""
    ds = [d.strip() for d in dates if d and d.strip()]
    if not ds:
        return ''
    if len(ds) == 1:
        return ds[0]
    if len(ds) == 2:
        return f"{ds[0]} si {ds[1]}"
    return ', '.join(ds[:-1]) + f" si {ds[-1]}"


def _build_pif_signature_table(doc, client_name, rep_client, rep_eg):
    """
    Tabel 1x2, fara chenar:
      Cell A: Beneficiar <CLIENT>  (bold)
              [3 linii goale]
              Ing. <rep_client>
      Cell B: Electroglobal SA  (bold)
              [3 linii goale]
              Ing. <rep_eg>
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    pairs = [
        (f"Beneficiar {client_name}".strip(), rep_client or ''),
        ("Electroglobal SA", rep_eg or 'Ing. Ion Ursu'),
    ]
    for cell_idx, (firma, rep) in enumerate(pairs):
        cell = table.rows[0].cells[cell_idx]
        cell.text = ''
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(firma)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        for _ in range(3):
            cell.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(rep)
        r2.font.size = Pt(10)
        r2.font.name = 'Calibri'
        _set_cell_borders(cell, sides=())
        _set_cell_width(cell, 8.0)
    return table


def _put_image_in_cell(cell, entry, pic_width_cm):
    """Pune poza + caption intr-o celula de tabel."""
    cell.text = ''
    p_pic = cell.paragraphs[0]
    p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_pic.add_run()
    try:
        run.add_picture(entry['file'], width=Cm(pic_width_cm))
    except Exception as e:
        run.add_text(f"[Imagine invalida: {e}]")
    caption = (entry.get('caption') or '').strip()
    if caption:
        p_cap = cell.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p_cap.add_run(caption)
        r2.italic = True
        r2.font.size = Pt(9)
        r2.font.name = 'Calibri'


def _build_images_block(doc, group, after_paragraph=None):
    """
    Construieste un tabel borderless cu pozele din `group`.
    Fiecare poza are `width`: 'full' (1 pe rand, 15cm) sau 'half' (2 pe rand, 7.5cm).
    Pozele 'half' consecutive se imperecheaza pe acelasi rand.
    `after_paragraph` — daca e dat, tabelul se insereaza dupa el.
    """
    n = len(group)
    if n == 0:
        return None

    # Construieste layout-ul randurilor respectand width per poza
    rows_layout = []
    i = 0
    while i < n:
        entry = group[i]
        w = (entry.get('width') or 'half').lower()
        if w == 'half' and i + 1 < n and (group[i + 1].get('width') or 'half').lower() == 'half':
            rows_layout.append([group[i], group[i + 1]])
            i += 2
        else:
            rows_layout.append([group[i]])
            i += 1

    table = doc.add_table(rows=len(rows_layout), cols=2)
    table.autofit = False
    cell_w = 8.0  # 16cm / 2

    for r, row_imgs in enumerate(rows_layout):
        if len(row_imgs) == 1:
            # poza singura — merge cele 2 celule, poza 15cm
            cell = table.rows[r].cells[0].merge(table.rows[r].cells[1])
            _put_image_in_cell(cell, row_imgs[0], pic_width_cm=15.0)
            _set_cell_borders(cell, sides=())
        else:
            for c in range(2):
                cell = table.rows[r].cells[c]
                _put_image_in_cell(cell, row_imgs[c], pic_width_cm=7.5)
                _set_cell_borders(cell, sides=())
                _set_cell_width(cell, cell_w)

    if after_paragraph is not None and after_paragraph._p is not None:
        anchor_p = after_paragraph._p
        parent = anchor_p.getparent()
        if parent is not None:
            idx = list(parent).index(anchor_p) + 1
            _move_table_to_position(table, parent, idx)
            spacer = OxmlElement('w:p')
            parent.insert(idx + 1, spacer)
    return table


def _group_images_ordered(images):
    """
    images: lista [{file, caption, anchor, width}].
    Returneaza lista de tuple (anchor, [imgs...]) pastrand ordinea de
    aparitie a ancorelor.
    """
    order = []
    buckets = {}
    for img in (images or []):
        a = img.get('anchor') or 'final'
        if a not in buckets:
            buckets[a] = []
            order.append(a)
        buckets[a].append(img)
    return [(a, buckets[a]) for a in order]


def _group_images_by_anchor(images):
    """images: lista de {file, caption, anchor}. Returneaza dict {anchor: [...]}"""
    by_anchor = {}
    for img in (images or []):
        a = img.get('anchor') or 'final'
        by_anchor.setdefault(a, []).append(img)
    return by_anchor


def generate_pv_pif(proiect, form_data):
    """
    proiect: dict din tabela proiecte (client, nume, nr_comanda, nr_contract, observatii, ...)
    form_data: dict cu cheile:
        cod_proiect, data_pv (DD.MM.YYYY),
        data_convocare, data_desfasurare (DD.MM.YYYY),
        date_probe: [DD.MM.YYYY, ...]
        nr_file (int, default 2), nr_anexe (int, default 0),
        constatari (str multi-linie, fiecare linie = paragraph),
        rep_client (str), rep_eg (str default 'Ing. Ion Ursu')
    """
    doc = Document(PIF_TEMPLATE)

    client = (proiect.get('client') or '').strip()
    obiectiv = (proiect.get('nume') or '').strip()
    nr_comanda = (proiect.get('nr_comanda') or '').strip()
    nr_contract = (proiect.get('nr_contract') or '').strip()
    comanda_contract = nr_contract or nr_comanda or ''
    if nr_comanda and nr_contract and nr_comanda != nr_contract:
        comanda_contract = f"{nr_comanda} / {nr_contract}"

    cod = form_data.get('cod_proiect', '').strip()
    data_pv = form_data.get('data_pv', datetime.now().strftime('%d.%m.%Y')).strip()
    data_conv = (form_data.get('data_convocare') or data_pv).strip()
    data_desf = (form_data.get('data_desfasurare') or data_conv).strip()
    date_probe = form_data.get('date_probe') or []
    probe_str = _format_date_list(date_probe) or data_desf
    nr_file = str(form_data.get('nr_file') or 2)
    nr_anexe = str(form_data.get('nr_anexe') or 0)
    constatari = (form_data.get('constatari') or '').strip()
    rep_client = (form_data.get('rep_client') or '').strip()
    rep_eg = (form_data.get('rep_eg') or 'Ing. Ion Ursu').strip()

    # Tabelul gol din header (1x2) — il las cum e
    # Iterez paragrafele si modific cele care matcheaza
    constatari_anchor = None  # paragraph dupa care vor fi inserate paragrafele de constatari

    for p in list(doc.paragraphs):
        txt = _para_text(p).strip()

        # Nr proces verbal
        if re.match(r'^Nr\.\s*\d+/.*data', txt, re.IGNORECASE):
            _set_paragraph_text(p, f"Nr. 1/{cod}/ data {data_pv}")
            continue
        # Beneficiar
        if txt.startswith('Beneficiar:'):
            _set_paragraph_text(p, f"Beneficiar: {client}")
            continue
        # Obiectivul de investitii
        if txt.startswith('Obiectivul de investi') or txt.startswith('Obiectivul de invest'):
            _set_paragraph_text(p, f"Obiectivul de investitii: {obiectiv}")
            continue
        # Comanda/contract
        if txt.startswith('Comanda/contract') or txt.startswith('Comanda /contract') or txt.startswith('Comanda / contract'):
            _set_paragraph_text(p, f"Comanda/contract nr: {comanda_contract}")
            continue
        # Comisia de receptie convocata...
        if 'Comisia de rece' in txt and 'convocat' in txt:
            _set_paragraph_text(
                p,
                f"Comisia de receptie convocata la data de {data_conv}: si-a desfasurat activitatea in data de {data_desf}."
            )
            continue
        # Punct 2: zilele cu probe tehnologice
        if txt.startswith('2.') and 'probele tehnologice' in txt:
            _set_paragraph_text(
                p,
                f"2. In zilele {probe_str} au fost efectuate probele tehnologice ale utilajelor si "
                f"instalatiilor aferente capacitatii pentru exploatarea normala a instalatiilor si "
                f"utilajelor tehnologice si asigurarea calitatii produselor, conform documentatiei "
                f"tehnico-economice."
            )
            continue
        # Punct 3: alte constatari (header)
        if txt.startswith('3.') and 'Alte constat' in txt:
            _set_paragraph_text(p, '3. Alte constatari:')
            constatari_anchor = p
            continue
        # Linii din vechile constatari (pana la "III. Concluzii" sau "Concluzii")
        if constatari_anchor is not None and txt:
            if txt.startswith('III.') or txt.startswith('Concluzii') or 'Pe baza constat' in txt:
                constatari_anchor = '__done__'  # marcheaza ca am terminat
                continue
            if constatari_anchor != '__done__':
                # sterge linia veche
                _delete_paragraph(p)
                continue

        # "Prezentul proces-verbal care contine X file si Y anexe..."
        if txt.startswith('Prezentul proces') and 'file' in txt:
            _set_paragraph_text(
                p,
                f"Prezentul proces-verbal care contine {nr_file} file si {nr_anexe} anexe numerotate, "
                f"cu un total de {nr_file} file care fac parte integranta din cuprinsul lui, a fost "
                f"incheiat astazi, {data_pv}, in 2 exemplare originale."
            )
            continue

    # Insereaza paragrafele de constatari dupa header-ul "3. Alte constatari:"
    if constatari and constatari != '__done__':
        # caut din nou anchor-ul (poate fi __done__ acum)
        for p in doc.paragraphs:
            if _para_text(p).strip().startswith('3.') and 'Alte constat' in _para_text(p):
                constatari_anchor = p
                break
    constatari_paras = []
    if True:
        if constatari_anchor and constatari_anchor != '__done__':
            from docx.text.paragraph import Paragraph as DocxParagraph
            anchor_p = constatari_anchor._p
            lines = [l.strip() for l in constatari.splitlines() if l.strip()]
            for line in lines:
                new_p = deepcopy(anchor_p)
                for t in new_p.iter(qn('w:t')):
                    t.text = ''
                anchor_p.addnext(new_p)
                wrap = DocxParagraph(new_p, constatari_anchor._parent)
                _set_paragraph_text(wrap, line)
                anchor_p = new_p
                constatari_paras.append(wrap)

    # ---- Poze pe ancore (PIF) ----
    # Anchor format: "constatari:N", "concluzii", "final"
    images = form_data.get('images') or []
    images_groups = _group_images_ordered(images)

    concluzii_para = None
    for p in doc.paragraphs:
        if 'Pe baza constat' in _para_text(p):
            concluzii_para = p
            break
    constatari_header = None
    for p in doc.paragraphs:
        if _para_text(p).strip().startswith('3.') and 'Alte constat' in _para_text(p):
            constatari_header = p

    pif_final_imgs = []
    for anchor_key, imgs in images_groups:
        if anchor_key.startswith('constatari:'):
            try:
                idx = int(anchor_key.split(':', 1)[1])
            except ValueError:
                idx = -1
            if 0 <= idx < len(constatari_paras):
                target = constatari_paras[idx]
            elif constatari_paras:
                target = constatari_paras[-1]
            else:
                target = constatari_header
            if target is not None:
                _build_images_block(doc, imgs, after_paragraph=target)
        elif anchor_key == 'concluzii':
            if concluzii_para is not None:
                _build_images_block(doc, imgs, after_paragraph=concluzii_para)
        elif anchor_key == 'final':
            pif_final_imgs.extend(imgs)

    # ---- Semnaturi: inlocuim tabelul existent cu unul nou ----
    old_sig = None
    for tbl in doc.tables:
        joined = ' '.join(c.text for row in tbl.rows for c in row.cells).lower()
        if 'beneficiar' in joined and 'electroglobal' in joined:
            old_sig = tbl
            break
    if old_sig is not None:
        parent, idx = _remove_table(old_sig)
        sig_table = _build_pif_signature_table(doc, client, rep_client, rep_eg)
        _move_table_to_position(sig_table, parent, idx)

    # Anexa final: pozele cu anchor "final"
    if pif_final_imgs:
        hdr_p = doc.add_paragraph()
        _add_page_break_before(hdr_p)
        run = hdr_p.add_run('Anexe foto')
        run.bold = True
        run.font.size = Pt(13)
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _build_images_block(doc, pif_final_imgs, after_paragraph=hdr_p)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
